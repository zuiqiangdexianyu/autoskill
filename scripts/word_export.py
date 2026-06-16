#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
word_export.py — 稿件导出 Word 文档

依赖：python-docx（已安装）
默认输出：桌面（自动检测路径）
格式：微软雅黑 12pt 纯黑色，标题居中加粗，小标题加粗居左，正文居左
"""

import os
import re
import shutil
import zipfile
import urllib.request
import tempfile
from datetime import datetime
from typing import List, Optional

# 当前会话的工作目录（同一对话内的过程文件存在同一目录）
_SESSION_WORK_DIR = None

# ── 常量 ──────────────────────────────────────────────────
# 字体名是纯字符串，与 docx 无关；尺寸/颜色常量依赖 docx 类型，放进 try 块，
# 避免未安装 python-docx 时模块在导入阶段就崩溃（实际导出时再报友好错误）。
FONT_NAME = "微软雅黑"

try:
    from docx import Document
    from docx.shared import Pt, Inches, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    HAS_DOCX = True
    FONT_SIZE = Pt(12)
    COLOR_BLACK = RGBColor(0x00, 0x00, 0x00)
    PAGE_WIDTH = Cm(16)  # A4 适中宽度
    MAX_IMG_WIDTH = Cm(14)
except ImportError:
    HAS_DOCX = False
    FONT_SIZE = COLOR_BLACK = PAGE_WIDTH = MAX_IMG_WIDTH = None


def get_desktop_path() -> str:
    """通过 PowerShell 获取桌面路径。"""
    import subprocess
    result = subprocess.check_output(
        ['powershell', '-NoProfile', '-Command', '[Environment]::GetFolderPath("Desktop")'],
        text=True, timeout=5
    ).strip()
    if result:
        return result
    return os.path.expanduser("~")


def get_work_dir() -> str:
    """获取本次会话的工作目录: SKILL安装目录/docx/YYYYMMDD_HHMMSS/

    同一会话内多次调用返回同一目录，仅实际写文件时创建。
    """
    global _SESSION_WORK_DIR
    if _SESSION_WORK_DIR is None:
        # SKILL 安装根目录 = scripts/../../
        skill_root = os.path.normpath(os.path.join(os.path.dirname(os.path.abspath(__file__)), ".."))
        ts = datetime.now().strftime("%Y%m%d_%H%M%S")
        _SESSION_WORK_DIR = os.path.join(skill_root, "docx", ts)
    return _SESSION_WORK_DIR


def clean_work_dir(work_dir: str):
    """清理临时工作目录。"""
    if os.path.isdir(work_dir):
        shutil.rmtree(work_dir, ignore_errors=True)


def save_process_md(content: str, filename: str) -> str:
    """保存过程文件（.md）到桌面工作目录，不会被自动清理。"""
    work_dir = get_work_dir()
    os.makedirs(work_dir, exist_ok=True)
    fname = filename if filename.endswith(".md") else filename + ".md"
    path = os.path.join(work_dir, fname)
    with open(path, "w", encoding="utf-8") as f:
        f.write(content)
    return path


def _set_run_font(run, size=FONT_SIZE, bold=False, color=COLOR_BLACK, name=FONT_NAME):
    """统一设置 run 的字体。"""
    run.font.name = name
    run.font.size = size
    run.font.color.rgb = color
    run.bold = bold
    run.element.rPr.rFonts.set(qn("w:eastAsia"), name)


def _add_formatted_paragraph(doc, text, bold=False, alignment=WD_ALIGN_PARAGRAPH.LEFT, font_size=FONT_SIZE, space_after=Pt(6)):
    """添加格式化段落。"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    _set_run_font(run, size=font_size, bold=bold)
    p.alignment = alignment
    p.paragraph_format.space_after = space_after
    p.paragraph_format.line_spacing = 1.5  # 1.5 倍行距
    return p


def _resolve_image(img_url: str) -> Optional[str]:
    """解析图片 URL/路径，返回本地可读文件路径。
    
    不会阻塞导出流程，失败直接返回 None。
    """
    if not img_url:
        return None
    # 1. 直接本地文件
    if os.path.isfile(img_url):
        return img_url
    # 2. 尝试 SKILL 图片库目录
    skill_root = os.path.normpath(os.path.join(os.path.dirname(os.path.abspath(__file__)), ".."))
    base, fname = os.path.dirname(img_url), os.path.basename(img_url)
    for c in [
        os.path.join(skill_root, "scripts", "images", fname),
        os.path.join(skill_root, "scripts", "images", "thumbs", fname),
        os.path.join(skill_root, "scripts", fname),
        img_url,  # 原路径再试一次（可能是相对路径）
    ]:
        if os.path.isfile(c):
            return c
    # 3. HTTP(S) 链接尝试下载（含 localhost——本地图片库 server 通常就跑在 localhost，
    #    在用户机器上是可达的；用超时避免 server 没起时卡住）
    if img_url.startswith(("http://", "https://")):
        try:
            ext = os.path.splitext(img_url)[1].split("?")[0] or ".jpg"
            tmp = tempfile.NamedTemporaryFile(delete=False, suffix=ext)
            with urllib.request.urlopen(img_url, timeout=10) as resp:
                tmp.write(resp.read())
            tmp.close()
            return tmp.name
        except Exception:
            return None
    return None


def _add_markdown_to_doc(doc, body):
    """将 Markdown 正文逐行解析写入 Word（纯 12pt 黑色微软雅黑）。"""
    lines = body.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i]

        if not line.strip():
            i += 1
            continue

        # H2 → 小标题，加粗居左
        if line.startswith("## "):
            text = line[3:].strip("* ")
            _add_formatted_paragraph(doc, text, bold=True, alignment=WD_ALIGN_PARAGRAPH.LEFT)
            i += 1
            continue

        # H3, H4 → 同 H2 处理
        if line.startswith("### "):
            text = line[4:].strip("* ")
            _add_formatted_paragraph(doc, text, bold=True, alignment=WD_ALIGN_PARAGRAPH.LEFT)
            i += 1
            continue
        if line.startswith("#### "):
            text = line[5:].strip("* ")
            _add_formatted_paragraph(doc, text, bold=True, alignment=WD_ALIGN_PARAGRAPH.LEFT)
            i += 1
            continue

        # 无序列表
        if re.match(r"^[\s]*[-*+]\s", line):
            items = []
            while i < len(lines) and re.match(r"^[\s]*[-*+]\s", lines[i]):
                items.append(re.sub(r"^[\s]*[-*+]\s", "", lines[i]))
                i += 1
            for item in items:
                _add_formatted_paragraph(doc, "• " + item)
            continue

        # 有序列表
        if re.match(r"^[\s]*\d+\.\s", line):
            idx = 1
            items = []
            while i < len(lines) and re.match(r"^[\s]*\d+\.\s", lines[i]):
                items.append(re.sub(r"^[\s]*\d+\.\s", "", lines[i]))
                i += 1
            for item in items:
                _add_formatted_paragraph(doc, f"{idx}. {item}")
                idx += 1
            continue

        # 引用 → 正文居左
        if line.startswith("> "):
            _add_formatted_paragraph(doc, line[2:])
            i += 1
            continue

        # 图片
        img_match = re.match(r"!\[.*?\]\((.*?)\)", line)
        if img_match:
            img_url = img_match.group(1)
            img_file = _resolve_image(img_url)
            if img_file:
                try:
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = p.add_run()
                    run.add_picture(img_file, width=MAX_IMG_WIDTH)
                except Exception:
                    pass
                finally:
                    try:
                        if img_file and os.path.isfile(img_file):
                            os.remove(img_file)
                    except Exception:
                        pass
            i += 1
            continue

        # 普通正文段落
        text = line.strip()
        if text:
            _add_formatted_paragraph(doc, text)
        i += 1


def save_article_as_docx(
    titles: List[str],
    body: str,
    output_path: str,
    style_name: str = "",
    topic: str = "",
) -> str:
    """将稿件保存为 Word 文档。

    文档格式：纯黑微软雅黑 12pt，标题居中加粗，小标题加粗居左，正文居左。
    """
    if not HAS_DOCX:
        raise RuntimeError("需要 python-docx 库: pip install python-docx")

    doc = Document()

    # 设置默认 Normal 样式
    norm = doc.styles["Normal"]
    norm.font.name = FONT_NAME
    norm.font.size = FONT_SIZE
    norm.font.color.rgb = COLOR_BLACK
    norm.element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)

    # 标题：居中加粗，每个标题后面空一行
    for t in titles:
        _add_formatted_paragraph(
            doc, t.strip(),
            bold=True,
            alignment=WD_ALIGN_PARAGRAPH.CENTER,
            space_after=Pt(4),
        )

    # 标题与正文之间空一行
    doc.add_paragraph()

    # 正文
    _add_markdown_to_doc(doc, body)

    # 保存
    out_dir = os.path.dirname(output_path)
    if out_dir:
        os.makedirs(out_dir, exist_ok=True)
    doc.save(output_path)
    return os.path.abspath(output_path)


def _safe_filename(text: str, max_len: int = 40) -> str:
    """清理非法文件名字符并截断。"""
    # 去掉非法字符：\/:*?"<>|
    safe = re.sub(r'[\\/:*?"<>|]', '', text)
    # 去掉首尾空白和点
    safe = safe.strip().strip('.')
    return safe[:max_len]


def check_titles_keyword(titles: List[str], keyword: str) -> List[str]:
    """返回不包含关键词的标题列表（用于强校验）。"""
    if not keyword:
        return []
    return [t for t in titles if keyword not in t]


def save_article_to_downloads(
    titles: List[str],
    body: str,
    style_name: str = "",
    topic: str = "",
    folder: str = "",
) -> str:
    """保存稿件到桌面（默认），文件名自动生成。

    导出前强校验：若传了 topic（核心关键词），3 个标题必须每个都包含它，
    否则抛出 ValueError——这是硬性要求，请把缺关键词的标题改了再导出。
    """
    missing = check_titles_keyword(titles, topic)
    if missing:
        raise ValueError(
            f"以下标题缺少核心关键词「{topic}」，必须每个标题都含该词，改完再导出：\n  - "
            + "\n  - ".join(missing)
        )
    ts = datetime.now().strftime("%Y%m%d_%H%M%S")
    safe_topic = re.sub(r'[\\/:*?"<>|]', "_", (topic or "稿件")[:20])
    fname = f"{safe_topic}_{ts}.docx"
    target = folder if folder else get_desktop_path()
    os.makedirs(target, exist_ok=True)
    output_path = os.path.join(target, fname)
    return save_article_as_docx(
        titles=titles, body=body,
        output_path=output_path,
        style_name=style_name, topic=topic,
    )


def save_batch_as_zip(
    articles: List[dict],
    output_path: str = "",
    style_name: str = "",
) -> str:
    """将多篇稿件打包为压缩包，保存到桌面（默认）。

    压缩包命名：主题-风格X篇-日期.zip
    每篇命名：序号、第一个标题.docx
    """
    if not articles:
        raise ValueError("稿件列表不能为空")

    topic = articles[0].get("topic", "稿件")
    count = len(articles)
    sn = style_name or articles[0].get("style_name", "")

    if not output_path:
        ts = datetime.now().strftime("%Y%m%d_%H%M%S")
        safe_topic = _safe_filename(topic, 20)
        zip_name = f"{safe_topic}-{sn}稿件{count}篇-{ts}.zip"
        output_path = os.path.join(get_desktop_path(), zip_name)

    docx_files = []
    work_dir = get_work_dir()
    os.makedirs(work_dir, exist_ok=True)
    docx_root = os.path.normpath(os.path.join(work_dir, ".."))  # docx/ 根目录
    try:
        for i, art in enumerate(articles):
            first_title = (art.get("titles") or [""])[0]
            safe_title = _safe_filename(first_title, 50)
            docx_name = f"{i+1}、{safe_title}.docx" if safe_title else f"{i+1}.docx"
            docx_path = os.path.join(work_dir, docx_name)
            save_article_as_docx(
                titles=art.get("titles", []),
                body=art.get("body", ""),
                output_path=docx_path,
                style_name=art.get("style_name", ""),
                topic=art.get("topic", ""),
            )
            docx_files.append(docx_path)

        with zipfile.ZipFile(output_path, "w", zipfile.ZIP_DEFLATED) as zf:
            for fp in docx_files:
                zf.write(fp, os.path.basename(fp))
        return os.path.abspath(output_path)
    finally:
        # 清理所有子目录（含子 Agent 创建的不同时间戳目录）
        if os.path.isdir(docx_root):
            for d in os.listdir(docx_root):
                p = os.path.join(docx_root, d)
                if os.path.isdir(p):
                    shutil.rmtree(p, ignore_errors=True)


if __name__ == "__main__":
    print("Word 导出工具加载成功。")
    print(f"  python-docx: {'可用' if HAS_DOCX else '未安装'}")
    print(f"  格式: 微软雅黑 {int(FONT_SIZE.pt)}pt 纯黑 | 标题居中加粗 | 小标题加粗居左 | 正文居左")
